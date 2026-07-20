from pathlib import Path
import json, re, time, urllib.parse, urllib.request
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT=Path(__file__).resolve().parents[1]
OUT=ROOT/"output"/"research"
CACHE=ROOT/"tmp"/"research"
OUT.mkdir(parents=True,exist_ok=True); CACHE.mkdir(parents=True,exist_ok=True)
DATA=json.loads((ROOT/"app/card-data.json").read_text(encoding="utf-8"))
SUBJECTS={"math":"数学","physics":"物理","nature":"自然科学","computer":"计算机科学"}
API="https://zh.wikipedia.org/w/api.php"
HEADERS={"User-Agent":"SciencePokerResearch/1.0 (educational local document builder)"}

VIDEO_LINKS={
 "大陆漂移说":"https://www.usgs.gov/educational-resources/usgs-educational-videos-and-animations",
 "水循环":"https://water.usgs.gov/vizlab/water-cycle/",
 "开普勒行星运动定律":"https://svs.gsfc.nasa.gov/4642/",
 "板块构造理论":"https://www.usgs.gov/educational-resources/usgs-educational-videos-and-animations",
 "光合作用总反应":"https://www.youtube.com/results?search_query=photosynthesis+animation+education",
 "DNA双螺旋与碱基配对":"https://www.youtube.com/results?search_query=DNA+double+helix+animation+education",
}

def wiki_search(title):
    params={"action":"query","generator":"search","gsrsearch":title,"gsrlimit":1,"prop":"extracts|pageimages|info","explaintext":1,"exsectionformat":"plain","piprop":"thumbnail|original","pithumbsize":900,"inprop":"url","format":"json","formatversion":2,"origin":"*"}
    try:
        url=API+"?"+urllib.parse.urlencode(params)
        req=urllib.request.Request(url,headers=HEADERS)
        data=json.loads(urllib.request.urlopen(req,timeout=25).read().decode("utf-8"))
        pages=data.get("query",{}).get("pages",[])
        if not pages:return None
        p=pages[0]; text=clean(p.get("extract", ""))
        return {"page_title":p.get("title",title),"text":text,"url":p.get("fullurl",f"https://zh.wikipedia.org/wiki/{urllib.parse.quote(title)}"),"image":(p.get("thumbnail") or p.get("original") or {}).get("source")}
    except Exception:return None

def clean(text):
    text=re.sub(r"\n{3,}","\n\n",text)
    text=re.sub(r"^(参见|参考资料|外部链接|注释|参考文献)\n[\s\S]*$","",text,flags=re.M)
    return text.strip()

def download_image(url,key):
    if not url:return None
    ext=".png" if ".png" in url.lower() else ".jpg"
    path=CACHE/f"{key}{ext}"
    if path.exists():return path
    try:
        req=urllib.request.Request(url,headers=HEADERS)
        content=urllib.request.urlopen(req,timeout=30).read(8_000_001)
        if len(content)>8_000_000:return None
        path.write_bytes(content); return path
    except Exception:return None

def set_font(run,size=11,bold=False,color="222222"):
    run.font.name="Microsoft YaHei"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei")
    run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)

def style_doc(doc,subject):
    sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.49)
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name="Microsoft YaHei"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor(35,43,52); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,color,before,after in [("Heading 1",16,"2E74B5",18,10),("Heading 2",13,"2E74B5",14,7),("Heading 3",12,"1F4D78",10,5)]:
        s=styles[name]; s.font.name="Microsoft YaHei"; s._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    styles["Heading 1"].font.size=Pt(18)
    header=sec.header.paragraphs[0]; header.text=f"自然科学文明扑克 · {SUBJECTS[subject]}资料集"; header.alignment=WD_ALIGN_PARAGRAPH.LEFT; set_font(header.runs[0],9,color="6B7785")
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); footer._p.append(fld)

def add_cover(doc,subject):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(120); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(f"{SUBJECTS[subject]}文明扑克"); set_font(r,30,True,"203748")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("54 张卡牌知识检索与素材记录"); set_font(r,16,False,"2E74B5")
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(60); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("内容包括：名词释义 · 原理与历史 · 应用与边界 · 图片素材 · 视频链接 · 检索来源"); set_font(r,10,color="6B7785")
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(140); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Science Poker Research Notes · 2026"); set_font(r,10,color="6B7785")
    doc.add_page_break()

def add_para(doc,text,bold_lead=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent=Inches(.28); p.paragraph_format.space_after=Pt(8); p.paragraph_format.line_spacing=1.35
    if bold_lead:
        r=p.add_run(bold_lead); set_font(r,10.5,True,"1F4D78")
    if text.startswith("http"):
        rid=p.part.relate_to(text,RT.HYPERLINK,is_external=True)
        hyperlink=OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"),rid)
        r=OxmlElement("w:r"); rpr=OxmlElement("w:rPr"); color=OxmlElement("w:color"); color.set(qn("w:val"),"2E74B5"); underline=OxmlElement("w:u"); underline.set(qn("w:val"),"single"); rpr.extend([color,underline]); t=OxmlElement("w:t"); t.text=text; r.extend([rpr,t]); hyperlink.append(r); p._p.append(hyperlink)
    else:
        r=p.add_run(text); set_font(r,10.5)
    return p

def shade_paragraph(paragraph,fill="F3F6F9"):
    ppr=paragraph._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); ppr.append(shd)
    paragraph.paragraph_format.left_indent=Inches(.12); paragraph.paragraph_format.right_indent=Inches(.12); paragraph.paragraph_format.space_before=Pt(4); paragraph.paragraph_format.space_after=Pt(12)

def add_formula_block(doc,formula):
    if not formula or formula.startswith("见牌面"): return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(12)
    shade_paragraph(p,"EEF4FA"); r=p.add_run(formula); set_font(r,12,False,"1F4D78"); r.font.name="Cambria Math"; r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei")

def section_text(card,wiki):
    title=card["title"]; formula=card.get("formula",""); impact=card.get("impact",""); credit=card.get("credit",""); suit=card.get("suit","")
    source=wiki["text"] if wiki else ""
    if len(source)>2400: source=source[:2400].rsplit("。",1)[0]+"。"
    overview=(source[:950].rsplit("。",1)[0]+"。") if len(source)>120 else f"{title}是{ suit }领域的重要概念。{impact}"
    remainder=source[len(overview):]
    principle=(remainder[:950].rsplit("。",1)[0]+"。") if len(remainder)>150 else f"理解{title}时，应先确认研究对象、关键变量、适用条件和结论之间的关系。牌面给出的公式或核心命题是知识的压缩表达，需要结合定义、单位、边界条件与具体情境阅读。"
    history=f"该知识在牌面中标注的历史或学术线索为“{credit}”。它被归入“{suit}”花色，说明其价值不仅在于一个结论，也在于它所代表的观察、建模、证明、实验或工程方法。{impact}"
    formula_text=(f"牌面表达：{formula}。" if formula and not formula.startswith("见牌面") else "牌面以核心命题或过程图景表达这一知识。")
    application=f"{formula_text}使用这一知识时，应区分理想模型与真实系统，核对量纲、定义域、初始条件、近似条件和测量误差。学习者还应比较它与同领域相邻概念的联系，避免只记住名称或公式而忽略适用范围。"
    return overview,principle,history,application

def add_card(doc,subject,index,card,wiki):
    rank=card["rank"]; title=card["title"]
    h=doc.add_paragraph(style="Heading 1"); h.paragraph_format.page_break_before=True; h.add_run(f"{index:02d}. {rank} · {title}")
    meta=doc.add_paragraph(); r=meta.add_run(f"花色：{card['suit']}    牌面来源：{card.get('credit','')} "); set_font(r,9,True,color="526579"); shade_paragraph(meta)
    add_formula_block(doc,card.get("formula",""))
    overview,principle,history,application=section_text(card,wiki)
    for heading,text in [("概述",overview),("核心原理与理解路径",principle),("历史、影响与学科位置",history),("应用、边界与学习提示",application)]:
        doc.add_paragraph(heading,style="Heading 2"); add_para(doc,text)
    image_path=download_image(wiki.get("image") if wiki else None,f"{subject}-{index:02d}")
    if image_path:
        doc.add_paragraph("图片素材",style="Heading 2")
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        try:p.add_run().add_picture(str(image_path),width=Inches(5.4))
        except Exception: pass
        cp=doc.add_paragraph(f"图：{wiki.get('page_title',title)}。来源见下方链接。"); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(cp.runs[0],8,color="6B7785")
    doc.add_paragraph("素材与来源",style="Heading 2")
    if wiki:add_para(doc,wiki["url"],"百科资料：")
    else:add_para(doc,"未检索到标题完全对应的百科条目，正文依据牌面原始说明整理。","检索说明：")
    if title in VIDEO_LINKS:add_para(doc,VIDEO_LINKS[title],"视频/动态素材：")
    add_para(doc,"本节为资料整理稿，适合继续用于网页编辑；涉及医学、工程安全或前沿争议时，应在发布前追加专业复核。","编辑提示：")

def build(subject):
    cache_file=CACHE/f"{subject}-research.json"
    research=json.loads(cache_file.read_text(encoding="utf-8")) if cache_file.exists() else []
    while len(research)<54:
        card=DATA[subject][len(research)]; print(subject,len(research)+1,card["title"],flush=True)
        research.append(wiki_search(card["title"])); cache_file.write_text(json.dumps(research,ensure_ascii=False,indent=2),encoding="utf-8"); time.sleep(.15)
    doc=Document(); style_doc(doc,subject); add_cover(doc,subject)
    doc.add_paragraph("使用说明",style="Heading 1"); add_para(doc,"本资料集按扑克牌顺序记录 54 个科学名词。每节以牌面原始信息为索引，结合中文百科检索结果整理，并保留图片、视频和来源链接。百科资料用于建立检索底稿，不等同于最终出版审校。")
    for i,(card,wiki) in enumerate(zip(DATA[subject],research),1): add_card(doc,subject,i,card,wiki)
    out=OUT/f"{SUBJECTS[subject]}文明扑克_54张知识与素材资料集.docx"; doc.save(out); return out

if __name__=="__main__":
    for s in SUBJECTS: print(build(s))
